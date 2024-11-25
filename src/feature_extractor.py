import csv
import os
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from win32com.client import Dispatch
import re
def rou(v):
    return round(v * 20) / 20

def extract_field_code(element):
    """Extracts a simplified representation of the field code from an XML element.

    Args:
        element: An lxml etree element representing a field.

    Returns:
        A string containing a simplified field code representation.
    """
    field_code_parts = []
    # Extract field code parts from the XML element
    for child in element.iterchildren():
        if child.tag == '{http://schemas.openxmlformats.org/wordprocessingml/2006/main}instrText':
            field_code_parts.append(child.text)
        elif child.tag == '{http://schemas.openxmlformats.org/wordprocessingml/2006/main}r':  # Handle simple text within the field
            field_code_parts.append(child.text)   

    return " ".join(field_code_parts)


def get_tab_positions(docx_paragraph, leftindent):  # Simplified parameters
    """Calculates tab positions, alignments, and indexes within a paragraph.

    Args:
        doc_path: Path to the Word document (.doc or .docx).
        para_number: The paragraph number (1-based).

    Returns:
        A tuple: (tab_info, tab_index)
            tab_info: A list of tuples, each containing: (tab position in cm, tab alignment)
            tab_index: A list of tab character indices within the paragraph text.
    """

    tab_info = []
    tab_index = []

    tab_stops = docx_paragraph.paragraph_format.tab_stops
    try:
        for tabstop in tab_stops: 
            position_cm = rou(tabstop.position/360680) 
            alignment = tabstop.alignment  # Extract alignment directly 
            tab_info.append((position_cm, alignment)) 
 
        tab_pattern = re.compile(r"\t") 
        for match in tab_pattern.finditer(docx_paragraph.text): 
            tab_index.append(match.start()) 

    except Exception as e:
        print(f"Error using docx: {e}. Add win32com fallback if needed.") 
        # ... (Potential win32com fallback - would need more specifics)

    return tab_info, tab_index 

    
word = Dispatch("Word.Application")
word.Visible = 0

# Open a document
folder_path = r"E:\Formatter\Misc\New features\Samples\Organizational Behavior FWU\English"

# Get a list of files in the folder
files = [f for f in os.listdir(folder_path) if os.path.isfile(os.path.join(folder_path, f)) and ((f.endswith(".docx") or f.endswith(".doc")) and not f.startswith("~"))]
print(files)
for f in files:
    input_doc_path = os.path.join(folder_path,f)
    if input_doc_path.split(".")[-1] == "docx":
        docx = True
        doc = word.Documents.Open(input_doc_path)
        document = Document(input_doc_path)
        document_width = document.sections[0].page_width
        document_left_margin = document.sections[0].left_margin
        document_right_margin = document.sections[0].right_margin
        margin_width = document_width - document_left_margin - document_right_margin
    else:
        docx = False
        doc = word.Documents.Open(input_doc_path)
        document = doc
        document_width = doc.sections[0].PageSetup.PageWidth
        document_left_margin = doc.sections[0].PageSetup.LeftMargin
        document_right_margin = doc.sections[0].PageSetup.RightMargin
        margin_width = document_width - document_left_margin - document_right_margin

    # Output CSV file
    output_csv_path = os.path.join(
        r"New features\latest\AfterDashain\OB",
        os.path.splitext(os.path.basename(input_doc_path))[0] + ".csv",
    )

    # Initialize CSV header and writer
    csv_header = (
        "Paragraph Number,Paragraph Style,First Line Indent,Hanging,Font Name,Font Size,Number of Characters,Tab Index,Number of Words,Number of Tabs,"
        "Number of Equals Signs,Italic,Bold,Paragraph Text,Numbering Type,Numbering Value,Tab Position,Within Table,Spacing Before,"
        "Spacing After,Table Width,Table Borders Style,Table Borders Color,Table Shading,Table Border Size,Table Font,"
        "Table Border Style,Table Indents"
    )

    with open(output_csv_path, "w", newline="", encoding='utf-8') as csv_file:
        writer = csv.writer(csv_file)
        writer.writerow(csv_header.split(","))

        # Loop through paragraphs
        para_number = 0
        para_range = doc.Range()
        prev_tab_count = []
        prev_tab_pos = []
        prev_para_style = []
        for para in doc.Paragraphs:
            para_range.Start = para.Range.Start
            para_range.End = para.Range.End
            para_range.Select()
            if para_range.Tables.Count <= 0:
                para_number += 1
            para_text = para_range.Text

            Fl_Indent = round(para.FirstLineIndent / 28.35, 2)
            Hanging = round(para.LeftIndent / 28.35, 2)
            char_count = para_range.Characters.Count
            word_count = len(para_range.Words)
            tab_count = para_range.Text.count("\t")
            eq_count = para_range.Text.count("=")
            is_italic = "Yes" if para_range.Italic else "No"
            is_bold = "Yes" if para_range.Bold else "No"
            font_name = para_range.Font.Name
            para_style = para.Style.NameLocal
            font_size = para_range.Font.Size
            prev_para_style.append(para_style)
            spacing_before = para.SpaceBefore
            spacing_after = para.SpaceAfter
            line_spacing = para_range.ParagraphFormat.LineSpacing
            numbering_type = para.Range.ListFormat.ListType
            docx_paragraph = document.paragraphs[para_number-1]
            if numbering_type != 0:
                list_value = para.Range.ListFormat.ListValue
            else:
                list_value = "0"

            if tab_count > 0:
                if docx:
                    pI = round(para.LeftIndent / 28.35, 2)
                    fI = round(para.FirstLineIndent / 28.35, 2)
                    tab_pos, tab_index = get_tab_positions(docx_paragraph, leftindent=pI)
                    # print(prev_tab_count,prev_tab_pos)

                else:
                    pI = round(para.LeftIndent / 28.35, 2)
                    fI = round(para.FirstLineIndent / 28.35, 2)
                    # tab_pos, tab_index = get_tab_positions(docx_paragraph, leftindent=pI)
            else:
                tab_pos = ""
                tab_index = ""

            paragraph = document.paragraphs[para_number - 1]  


            # para_text = "" 
            # field_map = {} 
            # field_marker_count = 0
            # current_field_text = "" 

            # for run in paragraph.runs:
            #     if run._element.xpath('.//w:fldChar[@w:fldCharType="begin"]'): 
            #         current_field_text = ""
            #         field_marker = f"{{FIELD_{field_marker_count}}}"
            #         para_text += field_marker
            #         field_map[field_marker] = "" 
                
            #     current_field_text += run.text  

            #     if run._element.xpath('.//w:fldChar[@w:fldCharType="end"]'):  
            #         field_map[field_marker] = extract_field_code(run._element)  
            #         field_marker_count += 1 
            #     else: 
            #         para_text += run.text  

            within_table = False
            table_width = ""
            table_borders_style = ""
            table_borders_color = ""
            table_shading = ""
            table_border_size = ""
            table_font = ""
            table_border_style = ""
            table_indents = ""
            para_text = para_text.replace(",","")
            if para_range.Tables.Count > 0 and within_table == False:
                within_table = True
                table = para_range.Tables(1)
                table_width = table.PreferredWidth
                table_borders_style = table.Borders.InsideLineStyle
                table_borders_color = table.Borders.InsideColor
                table_shading = table.Shading.BackgroundPatternColor
                table_border_size = table.Borders.InsideLineWidth
                table_font = table.Range.Font.Name
                table_border_style = table.Borders.InsideLineStyle
                if hasattr(table, "LeftIndent"):
                    table_indents = f"{table.LeftIndent},{table.RightIndent},{table.FirstColumnIndent}"

            try:
                writer.writerow(
                    [
                        para_number,
                        para_style,
                        Fl_Indent,
                        Hanging,
                        font_name,
                        font_size,
                        char_count,
                        tab_index,
                        word_count,
                        tab_count,
                        eq_count,
                        is_italic,
                        is_bold,
                        para_text,
                        numbering_type,
                        list_value,
                        tab_pos,
                        within_table,
                        spacing_before,
                        spacing_after,
                        table_width,
                        table_borders_style,
                        table_borders_color,
                        table_shading,
                        table_border_size,
                        table_font,
                        table_border_style,
                        table_indents,
                    ]
                )
            except Exception as e:
                print(f'Error: {str(e)}')
                continue

# Close the document and quit Word
    doc.Close(False)
